import docx #import the python-docx library to manipulate and edit docx using python
from docx.shared import RGBColor, Pt
show_text = docx.Document('show_text.docx') #the text to show
show_array = [] #to store the paragraphs in the text shown
for i in show_text.paragraphs:
	show_array.append(i.text)
hidden_text = docx.Document('hidden_text.docx') #the secret message
hidden_array = [] #to store the paragraphs in the hidden text
for i in hidden_text.paragraphs:
	if(len(i.text) != 0): #check for empty paragraphs
		hidden_array.append(i.text)
blueprint = docx.Document('blueprint.docx') #the blueprint file with all the predefined styles, etc.

def alter_spacing(para): #alter the line spacing between the paragraphs
	para_format = para.paragraph_format
	para_format.space_before = Pt(0)	
	para_format.space_after = Pt(0)	
	
no_of_hidden_lines = len(hidden_array)
current_hidden_line = 0 #a counter to check which line is being hidden right now
for i in show_array:
	if no_of_hidden_lines > current_hidden_line and i == "":
		para = blueprint.add_paragraph(hidden_array[current_hidden_line])
		para_counter = len(blueprint.paragraphs) - 1
		run = blueprint.paragraphs[para_counter].runs[0]
		font_style = run.font
		font_style.color.rgb = RGBColor(255, 255, 255) #font color set to white
		current_hidden_line += 1 #updating the hiddlen line counter
	else:
		para = blueprint.add_paragraph(i)
	alter_spacing(para)
blueprint.save('letterhead-cipher.docx')
print("Message Successfully Hidden")